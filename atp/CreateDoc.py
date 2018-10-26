from docx import Document
from docx.shared import Inches

'''
addXXX(content):
usage:
输入:相应内容string

addTablePara(txtPath, model_name, mma)
usage:
输入：路径string，模式名string，mma（max，min，ave）string
用于拼接为文件路径

addPicture(picturePath, model_name):
同上

saveDoc(filepath)
usage:
输入：路径string
'''


class CreateDoc(object):
    def __init__(self):
        self.__document = Document()
        self.head = self.__document.add_heading('ATP运行结果分析报告', 0)
        self.para_file = self.__document.add_paragraph(u'运行的ATP文件:')
        self.para_times = self.__document.add_paragraph(u'总仿真次数:')
        self.para_dataTable = self.__document.add_paragraph(u'运行所得数据图表:')

    def __addtoPara(self, paraNum, content):
        p = self.__document.paragraphs[paraNum]
        p.add_run('  ' + content)

    def addFilePara(self, content):
        self.__addtoPara(1, content)

    def addTimsPara(self, content):
        self.__addtoPara(2, content)

    def addTablePara(self, txtPath, mma):
        self.__document.add_paragraph(' ' + mma + '_table:')
        txtFileName = txtPath + '_%s.txt' % (mma)
        file = open(txtFileName)

        data = []
        for i in file.readlines():
            data.append(list(map(str.strip, i.split())))
        file.close()

        table = self.__document.add_table(len(data), len(data[0]))
        print(len(data), len(data[0]))
        i, j = 0, 0
        for row in table.rows:
            j = 0
            for cell in row.cells:
                cell.text = data[i][j]
                j += 1
            i += 1

    def addPicture(self, picturePath):
        pictureFileName = picturePath + '_pl4.png'
        self.__document.add_picture(pictureFileName, width=Inches(3))

    def saveDoc(self, filepath):
        self.__document.save(filepath + '.docx')


if __name__ == '__main__':
    doc = CreateDoc()
    doc.addFilePara('fuck')
    doc.addPicture(r'D:\dataatp\myfirst')
    doc.addTablePara(r'D:\dataatp\result', 'test')
    doc.saveDoc(r'D:\dataatp\myfirst')
